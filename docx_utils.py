# docx_utils.py
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Dict
import re

def extract_paragraphs(doc_path: str) -> List[str]:
    """
    Return a list of non-empty paragraph texts from a .docx file.
    """
    doc = Document(doc_path)
    return [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

def detect_headings(paragraphs: List[str]) -> str:
    """
    Simple heuristics to guess document type from the first ~30 paragraphs.
    """
    text = " ".join(paragraphs[:30]).lower()
    if "articles of association" in text or re.search(r'\b[aA][oO][aA]\b', text):
        return "Articles of Association"
    if "memorandum of association" in text or re.search(r'\b[mM][oO][aA]\b', text):
        return "Memorandum of Association"
    if "ubo declaration" in text or "ultimate beneficial owner" in text:
        return "UBO Declaration"
    # Add more heuristics as needed
    return "Unknown"

def annotate_and_save(original_path: str,
                      issues: List[Dict],
                      output_path: str):
    """
    Annotate the original .docx with:
      - yellow highlight on paragraphs that match issues
      - inline anchors like [ISSUE #n] appended to matched paragraphs
      - an appended "REVIEW ISSUES" section enumerating all issues
    issues is expected to be a list of dicts with keys:
      - 'para_index' (int) optional
      - 'match_text' (str) substring to find in the doc paragraph
      - 'issue_text' (str)
      - 'severity' (str)
      - 'suggestion' (str)
      - 'citation' (str)
    """
    doc = Document(original_path)

    # Map each issue to an index number for anchors
    # We'll try to match by substring match_text in paragraph text
    issue_counter = 1
    for iss in issues:
        match_text = iss.get("match_text", "")
        found = False
        if match_text:
            # Find the first paragraph that contains the match_text
            for p in doc.paragraphs:
                if match_text in p.text:
                    # Highlight all runs in that paragraph
                    # If no runs, create one (shouldn't normally happen)
                    if not p.runs:
                        run = p.add_run(p.text)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    else:
                        for run in p.runs:
                            try:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            except Exception:
                                # ignore runs that cannot be highlighted
                                pass
                    # Append anchor
                    p.add_run(f" [ISSUE #{issue_counter}]")
                    found = True
                    break
        if not found:
            # If no matching paragraph found, append a short note at the end
            appended = doc.add_paragraph(f"[ISSUE #{issue_counter}] (original text not found in body): {match_text[:200]}")
            # highlight note to make it visible
            for run in appended.runs:
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except Exception:
                    pass
        issue_counter += 1

    # Append issues appendix. Use safe fallback for Heading style absence.
    doc.add_page_break()
    try:
        # Preferred: use built-in heading style if available
        doc.add_heading("REVIEW ISSUES", level=1)
    except KeyError:
        # Fallback: create a bold, slightly larger paragraph to act as heading
        heading_para = doc.add_paragraph()
        run = heading_para.add_run("REVIEW ISSUES")
        run.bold = True
        run.font.size = Pt(14)
        heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Write out each issue in structured manner
    idx = 1
    for iss in issues:
        p = doc.add_paragraph()
        p.add_run(f"ISSUE #{idx}: ").bold = True
        p.add_run(iss.get("issue_text", "No issue text provided"))
        doc.add_paragraph(f"Matched text (excerpt): {iss.get('match_text', '')[:400]}")
        doc.add_paragraph(f"Severity: {iss.get('severity', 'Medium')}")
        doc.add_paragraph(f"Suggestion: {iss.get('suggestion', '')}")
        doc.add_paragraph(f"Citation: {iss.get('citation', '') or 'Pending — retrieve ADGM references'}")
        doc.add_paragraph("")  # spacer
        idx += 1

    # Save reviewed document
    doc.save(output_path)
